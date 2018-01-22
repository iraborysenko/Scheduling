from docx import Document
import general as g
from docx.shared import Inches


def build_report_first(sch, id):
    document = Document()

    head = document.add_heading('Звіт', 0)
    head.alignment = 1

    hd = document.add_heading('По виконавцю І-го рівня: ' + id[1].upper(), 1)
    hd.alignment = 1

    t = document.add_heading('Таблиця результатів', level=2)
    t.alignment = 1

    document.add_paragraph('На таблиці нижче наведено основну інформацію про завдання. ')

    document.add_paragraph('Час надходження робіт: ' + str(id[2]))

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Назва завдання'
    hdr_cells[1].text = 'Час початку виконання'
    hdr_cells[2].text = 'Час закінчення'
    hdr_cells[3].text = 'Директивний термін'
    hdr_cells[4].text = 'Термін виконання'

    for item in sch:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.title)
        row_cells[1].text = str(g.utc_to_str(item.start))
        row_cells[2].text = str(g.utc_to_str(item.start + item.length))
        row_cells[3].text = str(g.utc_to_str(item.deadline))
        row_cells[4].text = str(item.length/3600)

    document.add_paragraph("Максимальний старт: " + str(g.utc_to_str(sch[0].start)))

    document.add_page_break()

    v = document.add_heading('Візуалізація результатів', level=2)
    v.alignment = 1
    document.add_paragraph('На малюнку нижче наведено графік виконання робіт. ')
    document.add_picture('Pics/' + id[1] + '.png')
    document.add_paragraph('Різними кольорами позначено різні роботи. '
                           'Також тим самим кольором позначено відповідні директивні терміни.')

    document.save('Docs/' + id[1] + '.docx')


def build_report_second(sch, id):
    document = Document()

    head = document.add_heading('Звіт', 0)
    head.alignment = 1

    hd = document.add_heading('По виконавцю ІI-го рівня: ' + id[1].upper(), 1)
    hd.alignment = 1

    t = document.add_heading('Таблиця результатів', level=2)
    t.alignment = 1

    document.add_paragraph('На таблиці нижче наведено основну інформацію про завдання. ')

    document.add_paragraph('Час надходження робіт: ' + str(id[2]))

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Назва завдання'
    hdr_cells[1].text = 'Час початку виконання'
    hdr_cells[2].text = 'Час закінчення'
    hdr_cells[3].text = 'Директивний термін'
    hdr_cells[4].text = 'Термін виконання'
    hdr_cells[5].text = 'Ціна'

    criteria = 0
    totalsum = 0
    for item in sch:
        totalsum += int(item.price)
        criteria += item.deadline - item.start - item.length
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.title)
        row_cells[1].text = str(g.utc_to_str(item.start))
        row_cells[2].text = str(g.utc_to_str(item.start + item.length))
        row_cells[3].text = str(g.utc_to_str(item.deadline))
        row_cells[4].text = str(item.length/3600)
        row_cells[5].text = str(item.price)

    # document.add_paragraph("Значення критерію: " + str(int(criteria)))
    document.add_paragraph("Максимальний прибуток: " + str(totalsum))

    document.add_page_break()

    v = document.add_heading('Візуалізація результатів', level=2)
    v.alignment = 1
    document.add_paragraph('На малюнку нижче наведено графік виконання робіт. ')
    document.add_picture('Pics/' + id[1] + '.png')
    document.add_paragraph('Різними кольорами позначено різні роботи. '
                           'Також тим самим кольором позначено відповідні директивні терміни.')

    document.save('Docs/' + id[1] + '.docx')


def build_report_third(sch, excld, id):
    document = Document()

    head = document.add_heading('Звіт', 0)
    head.alignment = 1

    hd = document.add_heading('По виконавцю ІI-го рівня: ' + id[1].upper(), 1)
    hd.alignment = 1

    t = document.add_heading('Таблиця результатів', level=2)
    t.alignment = 1

    document.add_paragraph('На таблиці нижче наведено основну інформацію про завдання. ')

    document.add_paragraph('Час надходження робіт: ' + str(id[2]))

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Назва завдання'
    hdr_cells[1].text = 'Час початку виконання'
    hdr_cells[2].text = 'Час закінчення'
    hdr_cells[3].text = 'Директивний термін'
    hdr_cells[4].text = 'Термін виконання'
    hdr_cells[5].text = 'Ціна'

    totalsum = 0
    for item in sch:
        totalsum += int(item.price)
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.title)
        row_cells[1].text = str(g.utc_to_str(item.start))
        row_cells[2].text = str(g.utc_to_str(item.start + item.length))
        row_cells[3].text = str(g.utc_to_str(item.deadline))
        row_cells[4].text = str(item.length/3600)
        row_cells[5].text = str(item.price)

    document.add_paragraph('Масимальний прибуток: ' + str(totalsum))

    document.add_paragraph('')
    document.add_paragraph('Нижче наведено завдання, котрі були відкинуті. ')

    tableex = document.add_table(rows=1, cols=5)
    hdr_cells = tableex.rows[0].cells
    hdr_cells[0].text = 'Назва завдання'
    hdr_cells[1].text = 'Тривалість'
    hdr_cells[2].text = 'Ціна за раннє'
    hdr_cells[3].text = 'Ціна за пізнє'

    criteria = 0
    for item in excld:
        # criteria += item.deadline - item.start - item.length
        row_cells = tableex.add_row().cells
        row_cells[0].text = str(item[1])
        row_cells[1].text = str(item[4])
        row_cells[2].text = str(item[7])
        row_cells[3].text = str(item[8])


    document.add_page_break()

    v = document.add_heading('Візуалізація результатів', level=2)
    v.alignment = 1
    document.add_paragraph('На малюнку нижче наведено графік виконання робіт. ')
    document.add_picture('Pics/' + id[1] + '.png')
    document.add_paragraph('Різними кольорами позначено різні роботи. '
                           'Також тим самим кольором позначено відповідні директивні терміни.')

    document.save('Docs/' + id[1] + '.docx')